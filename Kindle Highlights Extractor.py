#pip install python-docx
import docx
import os
import pickle
import os.path
import sys
from googleapiclient.discovery import build
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
from googleapiclient.http import MediaFileUpload
from apiclient import errors
from apiclient.http import MediaFileUpload
from tkinter import Tk
from tkinter import filedialog

# If modifying these scopes, delete the file token.pickle.
SCOPES = ['https://www.googleapis.com/auth/drive.metadata.readonly',
          'https://www.googleapis.com/auth/drive.file']

#Enable Google Drive API and get credentials.json in the same folder as the program
def get_gdrive_service():
    creds = None
    # The file token.pickle stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    if os.path.exists('token.pickle'):
        with open('token.pickle', 'rb') as token:
            creds = pickle.load(token)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                'credentials.json', SCOPES)
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open('token.pickle', 'wb') as token:
            pickle.dump(creds, token)
    # return Google Drive API service
    return build('drive', 'v3', credentials=creds)


def main():
    master = Tk()
    master.withdraw()
    path1 = filedialog.askopenfilename(initialdir = "/",title = "Select My Clippings.txt file")
    while True:
        if str(path1).split('/')[-1] == 'My Clippings.txt':
            break
        elif str(path1) == '':
            sys.exit("You have not selected a file. The program is terminated.")
        else:
            print("Please select My Clippings.txt File only")
            path1 = filedialog.askopenfilename(initialdir = "/",title = "Please select My Clippings.txt File only")
    fo = open(str(path1),'r',encoding="utf8")
    mydoc = docx.Document()
    mydoc.add_heading("Highlights from my Books on Kindle", 0)
    mydoc.save("Highlights from Kindle.docx")

    L = fo.readlines()
    #To remove the newline list items
    L = list(filter(lambda a: a != '\n', L))

    D = {}
    for i in range(len(L)):
        try:
            if L[i] == '==========\n' and L[i+3] != "==========\n":            
                if str(L[i+1]) not in D:
                    D[str(L[i+1])] = [L[i+3]]
                else:
                    D[str(L[i+1])]+= [L[i+3]]
        except:
            #Exception raised when End of File is reached. So asking user to wait until all items of the file are saved.
            print('Please wait while we your highlights..')

    for i in D:
        tempL = []
        count = 0
        mydoc.add_heading(i, 1)
        mydoc.save("Highlights from Kindle.docx")
        for j in D[i]:
            if j not in tempL:
                count+=1
                tempL+=[j]
                mydoc.add_paragraph(str(count)+'. '+j)
                mydoc.save("Highlights from Kindle.docx")
        D[i] = tempL
    print('Your highlights are saved in the same folder as the program.')


def delete_file(service, file_id):
  try:
    service.files().delete(fileId=file_id).execute()
  except:
    print('An error occurred: %s' % error)


def upload_files():
    file_metadata = {
        "name": "Highlights from Kindle.docx",
    }
    media = MediaFileUpload("Highlights from Kindle.docx", resumable=True)
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    print("We have updated your highlights on your google drive.")



def findFile():
    path = 'Highlights from Kindle.docx'
    folder_query = "name = '%s'" % (path)
    folder_list = service.files().list(q=folder_query).execute()
    try:
        fileID = folder_list['files'][0]['id']
        delete_file(service,fileID)
    except:
        pass    
    

if __name__ == '__main__':
    service = get_gdrive_service()
    main()
    findFile()
    upload_files()